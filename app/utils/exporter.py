import io
from typing import Optional


def markdown_to_html(markdown_text: str) -> str:
    try:
        import markdown

        return markdown.markdown(markdown_text, extensions=["tables", "fenced_code"])
    except ImportError:
        lines = []
        for line in markdown_text.splitlines():
            if line.startswith("### "):
                lines.append(f"<h3>{line[4:].strip()}</h3>")
            elif line.startswith("## "):
                lines.append(f"<h2>{line[3:].strip()}</h2>")
            elif line.startswith("# "):
                lines.append(f"<h1>{line[2:].strip()}</h1>")
            else:
                lines.append(f"<p>{line}</p>")
        return "\n".join(lines)


def markdown_to_pdf(markdown_text: str) -> bytes:
    html = markdown_to_html(markdown_text)
    try:
        from weasyprint import HTML, CSS
    except ImportError as exc:
        raise RuntimeError("PDF export requires the weasyprint package") from exc

    css = CSS(string="""
        body { font-family: Arial, sans-serif; margin: 1rem; line-height: 1.5; }
        h1 { font-size: 28px; }
        h2 { font-size: 24px; }
        h3 { font-size: 20px; }
        p { font-size: 14px; margin-bottom: 0.75rem; }
    """)
    document = HTML(string=html)
    return document.write_pdf(stylesheets=[css])


def markdown_to_docx(markdown_text: str) -> bytes:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("DOCX export requires the python-docx package") from exc

    doc = DocxDocument()
    for line in markdown_text.splitlines():
        clean = line.strip()
        if not clean:
            doc.add_paragraph("")
            continue
        if clean.startswith("# "):
            doc.add_heading(clean[2:].strip(), level=1)
        elif clean.startswith("## "):
            doc.add_heading(clean[3:].strip(), level=2)
        elif clean.startswith("### "):
            doc.add_heading(clean[4:].strip(), level=3)
        else:
            doc.add_paragraph(clean)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_content_from_markdown(markdown_text: str, export_format: str) -> bytes:
    if export_format == "markdown":
        return markdown_text.encode("utf-8")
    if export_format == "pdf":
        return markdown_to_pdf(markdown_text)
    if export_format == "docx":
        return markdown_to_docx(markdown_text)
    raise ValueError(f"Unsupported export format: {export_format}")
